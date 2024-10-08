import re
import pdfplumber
from docx import Document
import pandas as pd
import streamlit as st
import io
import xlsxwriter

# Streamlit app title
st.title("CV Parser App")

# Function to extract all text from a PDF (adjusted for BytesIO file in Streamlit)
def extract_text_from_pdf(pdf_file):
    full_text = ""
    try:
        with pdfplumber.open(io.BytesIO(pdf_file.read())) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text(layout=True)
                if text:
                    full_text += text + "\n"
                else:
                    st.write(f"Warning: No text found on page {page_num + 1} of PDF file.")
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
    if not full_text.strip():
        st.write(f"No text was found in the PDF file.")
    return full_text

# Function to extract all text from a Word document (adjusted for BytesIO file in Streamlit)
def extract_text_from_word(docx_file):
    full_text = ""
    try:
        doc = Document(io.BytesIO(docx_file.read()))
        for para in doc.paragraphs:
            full_text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text + "\n"
    except Exception as e:
        st.error(f"Error extracting text from Word document: {e}")
    if not full_text.strip():
        st.write(f"No text was found in the Word document.")
    return full_text

# Enhanced Email Extraction (Handles emails split across lines or special formatting)
def extract_email(text):
    EMAIL_REGEX = r'[\w\.-]+@[\w\.-]+\.\w{2,3}'
    email_match = re.search(EMAIL_REGEX, text.replace("\n", ""))
    if email_match:
        email = email_match.group(0).replace(' ', '')
        return email

    lines = text.splitlines()
    for i, line in enumerate(lines):
        if "@" in line:
            before_at = re.search(r'[\w\.-]+', line)
            after_at = None
            if i + 1 < len(lines):
                after_at = re.search(r'[\w\.-]+\.\w{2,3}', lines[i + 1])
            if before_at and after_at:
                email = before_at.group(0) + "@" + after_at.group(0)
                return email.replace(' ', '')
    return None

# Phone Number Extraction
def extract_phone_number(text):
    PHONE_REGEX = r'(\+?\d{1,3}[- ]?\d{1,4}[- ]?\d{7,12})'
    phone_matches = re.findall(PHONE_REGEX, text)
    for phone in phone_matches:
        clean_phone = re.sub(r'[^+\d]', '', phone)
        if clean_phone.startswith('+') and len(clean_phone) == 12:
            return clean_phone
        elif clean_phone.startswith('0') and len(clean_phone) == 11:
            return clean_phone
    return None

# Function to extract Name (Focuses on the largest font size and first few words)
def extract_name(extracted_text):
    lines = extracted_text.splitlines()
    IGNORE_KEYWORDS = ["PROFILE", "SKILLS", "EXPERIENCE", "CONTACT", "EDUCATION", "OBJECTIVE", "SUMMARY"]

    potential_name = None
    for line in lines[:10]:  # Scan first 10 lines for the name
        line = line.strip()
        if not line or any(keyword in line.upper() for keyword in IGNORE_KEYWORDS):
            continue
        if re.match(r"^[A-Za-z\s]+$", line) and len(line.split()) <= 4:
            potential_name = line
            break
    return potential_name

# Function to clean and extract name, email, and phone from the extracted text
def extract_info_from_text(extracted_text):
    info = {
        "Name": extract_name(extracted_text),
        "Email": extract_email(extracted_text),
        "Phone": extract_phone_number(extracted_text)
    }
    return info

# Function to process files (PDF, Word)
def process_file(file, file_type):
    extracted_text = ""
    if file_type == '.pdf':
        extracted_text = extract_text_from_pdf(file)
    elif file_type == '.docx':
        extracted_text = extract_text_from_word(file)
    return extract_info_from_text(extracted_text)

# Main Streamlit logic
uploaded_files = st.file_uploader("Upload CVs (PDF or Word)", type=['pdf', 'docx'], accept_multiple_files=True)

if uploaded_files:
    data = []
    for uploaded_file in uploaded_files:
        file_name = uploaded_file.name
        file_extension = file_name.split('.')[-1].lower()

        st.write(f"Processing {file_name}...")

        info = process_file(uploaded_file, f'.{file_extension}')
        info['File Name'] = file_name
        data.append(info)

    # Create DataFrame from extracted data
    df = pd.DataFrame(data)

    # Display extracted data in the app
    st.dataframe(df)

    # Create a BytesIO buffer to save the Excel file in memory
    output = io.BytesIO()

    # Write the data to the Excel file using ExcelWriter and xlsxwriter
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False)
        writer.close()

    # Move the buffer position to the beginning
    output.seek(0)

    # Provide download button for the extracted data as an Excel file
    st.download_button(
        label="Download extracted data as Excel",
        data=output,
        file_name="extracted_cv_data.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
